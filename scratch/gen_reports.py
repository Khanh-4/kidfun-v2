#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Tạo 2 file báo cáo tuần Sprint 7 cho dự án KidFun V3."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure East Asian font matches
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.append(rFonts)

def add_heading(doc, text, level=1, align=WD_ALIGN_PARAGRAPH.LEFT, size=None, bold=True, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 13}
    s = size or sizes.get(level, 12)
    set_font(run, size=s, bold=bold, color=color)
    return p

def add_para(doc, text, indent=0, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_table_row(table, values, bold_first=False, bg=None, sizes=None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        s = sizes[i] if sizes else 11
        is_bold = bold_first and i == 0
        set_font(run, size=s, bold=is_bold)
        if bg:
            set_cell_bg(cell, bg)
    return row

def make_header_row(table, headers, bg='4472C4', text_color=(255,255,255)):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=11, bold=True, color=text_color)
        set_cell_bg(cell, bg)


# ═══════════════════════════════════════════════════════════════════
# FILE 1: BÁO CÁO CHI TIẾT TIẾN ĐỘ TUẦN (Sprint 7)
# ═══════════════════════════════════════════════════════════════════

def create_detailed_report():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    # ── TIÊU ĐỀ ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HCM (HUTECH)')
    set_font(run, size=13, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
    set_font(run, size=13, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BÁO CÁO TIẾN ĐỘ TUẦN — SPRINT 7')
    set_font(run, size=16, bold=True, color=(31, 73, 125))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ĐỒ ÁN CƠ SỞ NGÀNH CÔNG NGHỆ PHẦN MỀM')
    set_font(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Đề tài: KidFun V3 — Ứng dụng kiểm soát thời gian sử dụng thiết bị của trẻ em')
    set_font(run, size=12, italic=True)

    doc.add_paragraph()

    # Thông tin chung (bảng)
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ('Nhóm', 'Nhóm 60 — 2 thành viên'),
        ('Thành viên', 'Khanh (Backend) | Arix (Frontend / Mobile)'),
        ('Sprint', 'Sprint 7 — GPS, Geofencing & SOS'),
        ('Thời gian', '07/04/2026 – 12/04/2026 (1 tuần)'),
        ('Tuần', 'Tuần 7/10 của kế hoạch 10 sprint'),
        ('Ngày báo cáo', '12/04/2026'),
    ]
    for i, (k, v) in enumerate(info):
        row = tbl.rows[i]
        for j, txt in enumerate([k, v]):
            cell = row.cells[j]
            p2 = cell.paragraphs[0]
            run = p2.add_run(txt)
            set_font(run, size=11, bold=(j == 0))
            if j == 0:
                set_cell_bg(cell, 'D9E2F3')
            cell.width = Cm(4.5) if j == 0 else Cm(12)

    doc.add_paragraph()

    # ── 1. TỔNG QUAN SPRINT 7 ────────────────────────────────────
    add_heading(doc, 'I. TỔNG QUAN SPRINT 7', level=1, color=(31, 73, 125))

    add_para(doc,
        'Sprint 7 của dự án KidFun V3 diễn ra từ ngày 07/04/2026 đến 12/04/2026, '
        'tập trung vào việc triển khai nhóm tính năng vị trí và an toàn cho trẻ em '
        'gồm ba module chính: (1) GPS Tracking thời gian thực, (2) Geofencing — vùng '
        'an toàn tự động cảnh báo, và (3) Nút SOS khẩn cấp kèm ghi âm. Đây là một '
        'trong những sprint quan trọng nhất trong giai đoạn hoàn thiện sản phẩm trước '
        'khi bảo vệ trước hội đồng.',
        first_line=True, size=12)

    add_para(doc,
        'Mục tiêu tổng thể của sprint đã được hoàn thành 100%: toàn bộ 24 test case '
        'trong kế hoạch kiểm thử đều pass, không có lỗi nào tồn đọng. Nhóm đã thực '
        'hiện tổng cộng 107 commit và xử lý 26 Pull Request trong vòng 6 ngày làm '
        'việc, đây là sprint có mật độ code cao nhất kể từ đầu dự án.',
        first_line=True, size=12)

    # Bảng tổng quan
    add_heading(doc, '1.1. Kết quả tổng hợp', level=2, size=13)
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = 'Table Grid'
    make_header_row(tbl2, ['Module', 'Tính năng', 'Số TC', 'Kết quả'])
    data2 = [
        ('GPS Tracking', 'Theo dõi vị trí real-time, lịch sử di chuyển', '5', '5/5 Pass ✓'),
        ('Geofencing', 'Tạo/xóa vùng an toàn, ENTER/EXIT detection', '7', '7/7 Pass ✓'),
        ('SOS', 'Nút SOS khẩn cấp, ghi âm 15s, push notification', '9', '9/9 Pass ✓'),
        ('API Validation', 'Auth guard, error handling, upload validation', '3', '3/3 Pass ✓'),
        ('TỔNG', '', '24', '24/24 Pass ✓'),
    ]
    for i, row_data in enumerate(data2):
        r = add_table_row(tbl2, row_data, bg='E8F4E8' if i % 2 == 0 else None)
        if i == 4:
            for cell in r.cells:
                set_cell_bg(cell, 'C6EFCE')
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True

    doc.add_paragraph()

    # ── 2. BACKEND ────────────────────────────────────────────────
    add_heading(doc, 'II. PHẦN BACKEND (Thành viên: Khanh)', level=1, color=(31, 73, 125))

    add_para(doc,
        'Phần backend do thành viên Khanh phụ trách, bao gồm thiết kế database, '
        'xây dựng REST API, xử lý logic Geofence, và tích hợp push notification. '
        'Tất cả các task đều được hoàn thành trong đúng tiến độ kế hoạch.',
        first_line=True)

    # 2.1 Database
    add_heading(doc, '2.1. Database Models (Task 1)', level=2, size=13)
    add_para(doc,
        'Đã bổ sung 4 model mới vào Prisma schema và thực hiện migration thành công '
        'lên PostgreSQL Railway:',
        first_line=True)

    bullets_db = [
        'LocationLog — lưu tọa độ GPS theo thời gian (profileId, deviceId, latitude, longitude, accuracy, source, createdAt). Index theo [profileId, createdAt] để query nhanh.',
        'Geofence — vùng an toàn do phụ huynh tạo (name, latitude, longitude, radius 50–5000m, isActive). Liên kết với Profile qua Cascade delete.',
        'GeofenceEvent — lịch sử sự kiện ENTER/EXIT (geofenceId, profileId, type, latitude, longitude, createdAt). Index theo [profileId, createdAt].',
        'SOSAlert — cảnh báo khẩn cấp (profileId, deviceId, vị trí, audioUrl, message, status ACTIVE/ACKNOWLEDGED/RESOLVED, acknowledgedAt, resolvedAt).',
    ]
    for b in bullets_db:
        add_bullet(doc, b)

    # 2.2 Location API
    add_heading(doc, '2.2. Location Tracking API (Task 2)', level=2, size=13)
    add_para(doc,
        'Đã xây dựng LocationController với 3 endpoint chính:',
        first_line=True)
    add_bullet(doc, 'POST /api/child/location — Child gửi GPS (không cần auth, chỉ cần deviceCode). Sau khi lưu log, hệ thống emit Socket.IO event "locationUpdated" đến phòng family_{userId} và gọi geofence check.')
    add_bullet(doc, 'GET /api/profiles/:id/location/current — Phụ huynh lấy vị trí mới nhất của con.')
    add_bullet(doc, 'GET /api/profiles/:id/location/history?date=YYYY-MM-DD — Lịch sử vị trí theo ngày, trả về mảng có thứ tự tăng dần theo thời gian để vẽ polyline.')

    # 2.3 Geofence
    add_heading(doc, '2.3. Geofence CRUD API và ENTER/EXIT Detection (Task 3 & 4)', level=2, size=13)
    add_para(doc,
        'Đã xây dựng GeofenceController với đầy đủ CRUD và GeofenceService xử lý '
        'logic phát hiện ENTER/EXIT dựa trên công thức Haversine:',
        first_line=True)
    add_bullet(doc, 'GET /api/profiles/:id/geofences — Danh sách tất cả vùng an toàn.')
    add_bullet(doc, 'POST /api/profiles/:id/geofences — Tạo mới (validate radius 50–5000m).')
    add_bullet(doc, 'PUT /api/geofences/:id — Cập nhật thông tin hoặc bật/tắt (isActive).')
    add_bullet(doc, 'DELETE /api/geofences/:id — Xóa (cascade xóa cả sự kiện liên quan).')
    add_bullet(doc, 'GET /api/profiles/:id/geofences/events?date=YYYY-MM-DD — Lịch sử ENTER/EXIT.')
    add_para(doc,
        'GeofenceService sử dụng in-memory cache Map<"profileId_geofenceId", boolean> '
        'để theo dõi trạng thái trong/ngoài. Mỗi khi child gửi GPS, hệ thống tự động '
        'tính khoảng cách Haversine đến tất cả geofence active của profile đó và sinh '
        'GeofenceEvent khi trạng thái thay đổi.',
        first_line=True)

    # 2.4 SOS
    add_heading(doc, '2.4. SOS Alert API với Audio Upload (Task 5)', level=2, size=13)
    add_para(doc,
        'Đã tích hợp package Multer cho file upload và xây dựng SOSController:',
        first_line=True)
    add_bullet(doc, 'POST /api/child/sos (multipart/form-data) — Child gửi SOS kèm file âm thanh (.m4a, .mp3, .aac, .wav, tối đa 5MB). Hệ thống lưu file vào /uploads/sos-audio/, tạo SOSAlert record, emit Socket.IO "sosAlert" và gửi FCM push notification.')
    add_bullet(doc, 'GET /api/profiles/:id/sos — Lịch sử SOS 50 bản ghi gần nhất.')
    add_bullet(doc, 'PUT /api/sos/:id/acknowledge — Phụ huynh xác nhận (ACKNOWLEDGED + acknowledgedAt).')
    add_bullet(doc, 'PUT /api/sos/:id/resolve — Phụ huynh đánh dấu giải quyết xong (RESOLVED + resolvedAt).')
    add_para(doc,
        'Lưu ý: Multer error handling được bổ sung sau khi phát hiện bug TC-24 — '
        'lỗi upload file sai định dạng trả về 500 thay vì 400. Đã fix bằng cách '
        'wrap middleware trong try-catch và xử lý MulterError riêng biệt.',
        first_line=True, italic=True)

    # 2.5 Push Notification
    add_heading(doc, '2.5. Push Notifications (Task 7)', level=2, size=13)
    add_para(doc,
        'Đã mở rộng FCMService với 2 function mới:',
        first_line=True)
    add_bullet(doc, 'sendGeofencePushNotification() — Gửi FCM notification khi detect ENTER/EXIT. Title tùy theo loại sự kiện (VD: "Bé Nam đã vào Trường học"), body mô tả trạng thái an toàn. Tự động xóa stale token khi gặp lỗi "registration-token-not-registered".')
    add_bullet(doc, 'sendSOSPushNotification() — Gửi FCM với priority MAX, channel "sos_critical", sound default. Đảm bảo phụ huynh nhận thông báo ngay cả khi app đang tắt hoàn toàn.')

    # 2.6 PR list
    add_heading(doc, '2.6. Danh sách Pull Request Backend (Khanh)', level=2, size=13)
    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = 'Table Grid'
    make_header_row(tbl3, ['PR#', 'Branch', 'Mô tả', 'Ngày'])
    pr_backend = [
        ('#143', 'feature/backend/location-models', 'Thêm 4 Prisma models: Location, Geofence, SOS', '08/04'),
        ('#144', 'feature/backend/location-api', 'Location Tracking API (POST + GET current/history)', '08/04'),
        ('#145', 'feature/backend/geofence-crud', 'Geofence CRUD API + ENTER/EXIT detection', '08/04'),
        ('#146', 'feature/backend/sos-alert', 'SOS Alert API với Multer audio upload', '08/04'),
        ('#147', 'feature/backend/location-push-notifications', 'FCM push notification cho Geofence + SOS', '08/04'),
        ('#157', 'chore/update-package-lock', 'Cập nhật package-lock.json sau khi thêm multer', '09/04'),
        ('#187', 'fix/mobile/sprint7-notification-bugs', 'Fix duplicate SOS notification + geofence tên mặc định', '11/04'),
        ('#188', 'fix/backend/tc24-multer-error-handling', 'Fix Multer error handling (trả 400 thay vì 500)', '11/04'),
    ]
    for i, row_data in enumerate(pr_backend):
        add_table_row(tbl3, row_data, bg='EBF3FB' if i % 2 == 0 else None)

    doc.add_paragraph()

    # ── 3. FRONTEND / MOBILE ──────────────────────────────────────
    add_heading(doc, 'III. PHẦN MOBILE / FRONTEND (Thành viên: Arix)', level=1, color=(31, 73, 125))

    add_para(doc,
        'Phần frontend mobile do thành viên Arix phụ trách, xây dựng trên nền '
        'Flutter 3.x với Mapbox SDK cho bản đồ. Sprint 7 là sprint đầu tiên sử dụng '
        'bản đồ nên phải trải qua nhiều vòng fix build error liên quan đến Mapbox '
        'trước khi chạy được trên thiết bị thật.',
        first_line=True)

    # 3.1 Mapbox
    add_heading(doc, '3.1. Mapbox SDK Setup (Task 1)', level=2, size=13)
    add_para(doc,
        'Đã cài đặt và cấu hình thành công mapbox_maps_flutter. Các công việc chính:',
        first_line=True)
    add_bullet(doc, 'Thêm Maven repository của Mapbox vào build.gradle với xác thực MAPBOX_DOWNLOADS_TOKEN.')
    add_bullet(doc, 'Cấu hình permissions trong AndroidManifest.xml: ACCESS_FINE_LOCATION, ACCESS_COARSE_LOCATION, ACCESS_BACKGROUND_LOCATION, RECORD_AUDIO.')
    add_bullet(doc, 'Khởi tạo MapboxOptions.setAccessToken() trong main.dart, load token từ file .env (không commit vào git).')
    add_bullet(doc, 'Giải quyết xung đột namespace: class Size trong Flutter và Mapbox — prefix import mapbox. Đây là bug phức tạp nhất liên quan đến build, mất nhiều thời gian debug.')
    add_bullet(doc, 'Fix blank map bằng cách truyền styleUri và textureView vào MapWidget.')

    # 3.2 GPS Service
    add_heading(doc, '3.2. GPS Tracking Service (Task 2 & 3)', level=2, size=13)
    add_para(doc,
        'Đã xây dựng LocationService singleton với interval linh hoạt:',
        first_line=True)
    add_bullet(doc, 'Foreground mode: gửi vị trí mỗi 30 giây.')
    add_bullet(doc, 'Background mode: gửi vị trí mỗi 5 phút (tiết kiệm pin).')
    add_bullet(doc, 'WidgetsBindingObserver theo dõi lifecycle để tự động chuyển đổi interval.')
    add_bullet(doc, 'Child App tự động sync location lên backend qua LocationRepository sau mỗi update.')
    add_bullet(doc, 'Fix bug đếm ngược thời gian kẹt tại 0 (delta logic lỗi khi parent cập nhật time limit).')

    # 3.3 Map Screen
    add_heading(doc, '3.3. Parent App — Map Screen (Task 4)', level=2, size=13)
    add_para(doc,
        'Màn hình bản đồ phụ huynh hiển thị vị trí trẻ em real-time:',
        first_line=True)
    add_bullet(doc, 'MapWidget với PointAnnotationManager hiển thị marker vị trí con.')
    add_bullet(doc, 'Lắng nghe Socket.IO event "locationUpdated" để tự động dịch chuyển marker không cần reload.')
    add_bullet(doc, 'FAB "Lấy vị trí hiện tại" gọi API GET location/current và center map vào zoom 15.')
    add_bullet(doc, 'Fix bug marker không hiển thị do thiếu asset icon (cần load từ Flutter assets).')

    # 3.4 Geofence UI
    add_heading(doc, '3.4. Parent App — Geofence UI (Task 5)', level=2, size=13)
    add_para(doc,
        'Giao diện quản lý vùng an toàn trực quan trên bản đồ:',
        first_line=True)
    add_bullet(doc, 'Chế độ thêm mới: tap lên bản đồ để chọn tâm geofence, slider điều chỉnh radius (50–5000m), vẽ CircleAnnotation theo radius thay đổi real-time.')
    add_bullet(doc, 'Dialog nhập tên vùng (Nhà, Trường học, v.v.) trước khi lưu.')
    add_bullet(doc, 'Chế độ xóa: tap vào polygon geofence hiện AlertDialog xác nhận xóa.')
    add_bullet(doc, 'Fix UI: 2 nút trong dialog (Lưu / Hủy) có kích thước bằng nhau — dùng Row + Expanded.')
    add_bullet(doc, 'Lắng nghe Socket.IO "geofenceEvent" để hiện dialog thông báo ENTER/EXIT real-time.')
    add_bullet(doc, 'Fix race condition: nếu map chưa load xong mà đã nhận geofence event thì bỏ qua thay vì crash.')

    # 3.5 Location History
    add_heading(doc, '3.5. Parent App — Lịch sử Vị trí (Task 6)', level=2, size=13)
    add_para(doc,
        'Màn hình xem lại hành trình di chuyển trong ngày:',
        first_line=True)
    add_bullet(doc, 'DatePicker chọn ngày → gọi API GET location/history?date=.')
    add_bullet(doc, 'Vẽ LineLayer (polyline) trên Mapbox nối các điểm GPS theo thứ tự thời gian.')
    add_bullet(doc, 'Fix bug LineLayer.lineColor — phải dùng CSS color string ("rgba(0,0,255,1)") thay vì int hex.')
    add_bullet(doc, 'Fix bug cameraForCoordinateBounds khi chỉ có 1 điểm.')
    add_bullet(doc, 'Tích hợp hiển thị cả location logs và geofence events trong cùng màn hình lịch sử (TC-12).')
    add_bullet(doc, 'Auto-refresh geofence events real-time khi nhận Socket.IO event thay vì phải chọn lại ngày.')

    # 3.6 SOS
    add_heading(doc, '3.6. Child App — Nút SOS (Task 7)', level=2, size=13)
    add_para(doc,
        'Nút SOS khẩn cấp trên giao diện trẻ em:',
        first_line=True)
    add_bullet(doc, 'FloatingActionButton màu đỏ nổi bật với icon SOS.')
    add_bullet(doc, 'Confirm dialog "🆘 Gửi SOS?" để tránh gửi nhầm.')
    add_bullet(doc, 'Gửi SOS fast alert ngay lập tức (< 2s), sau đó ghi âm 15s và upload.')
    add_bullet(doc, 'Fix: SOS dialog không hiển thị trên Child app (chỉ hiển thị trên Parent app).')
    add_bullet(doc, 'Fix: cấp quyền microphone đúng luồng.')

    # 3.7 Parent SOS
    add_heading(doc, '3.7. Parent App — SOS Alert Screen (Task 8)', level=2, size=13)
    add_para(doc,
        'Màn hình nhận và xử lý cảnh báo khẩn cấp:',
        first_line=True)
    add_bullet(doc, 'SOS Alert Dialog với barrierDismissible = false (bắt buộc phụ huynh phải tương tác).')
    add_bullet(doc, 'Nút "Nghe ghi âm" (AudioPlayer) xuất hiện khi có audioUrl.')
    add_bullet(doc, 'Nút "Xem vị trí" navigate sang Map Screen centered vào SOS location.')
    add_bullet(doc, 'Nút "Gọi con" mở Dialer với số điện thoại của trẻ.')
    add_bullet(doc, 'Nút "Đã nhận được" gọi API acknowledge → đổi status.')
    add_bullet(doc, 'SOS History Screen với status badge (ACTIVE = đỏ, ACKNOWLEDGED = vàng, RESOLVED = xanh).')
    add_bullet(doc, 'Fix double UI: notification tap không mở 2 dialog cùng lúc.')
    add_bullet(doc, 'Fix FCM navigation: tap notification khi app killed sẽ mở đúng màn hình SOS.')
    add_bullet(doc, 'Fix SOS history 404 khi chưa có data.')

    # 3.8 Notifications
    add_heading(doc, '3.8. Flutter Local Notifications (Background Alerts)', level=2, size=13)
    add_para(doc,
        'Đã tích hợp flutter_local_notifications để hiển thị thông báo khi app chạy '
        'ngầm hoặc bị kill. Phải enable coreLibraryDesugaring trong build.gradle do '
        'flutter_local_notifications yêu cầu Java 8+ desugaring.',
        first_line=True)

    # PR list Mobile
    add_heading(doc, '3.9. Danh sách Pull Request Mobile (Arix)', level=2, size=13)
    tbl4 = doc.add_table(rows=1, cols=4)
    tbl4.style = 'Table Grid'
    make_header_row(tbl4, ['PR#', 'Branch', 'Mô tả', 'Ngày'])
    pr_mobile = [
        ('#150', 'feature/mobile/mapbox-setup', 'Setup Mapbox SDK + permissions', '08/04'),
        ('#148', 'feature/mobile/location-history', 'Parent: Location History Screen', '08/04'),
        ('#149', 'feature/mobile/child-sos', 'Child: SOS button + audio recording', '08/04'),
        ('#153', 'feature/mobile/parent-sos', 'Parent: SOS Alert dialog + history', '08/04'),
        ('#155', 'feature/mobile/gps-tracking', 'Sprint 7 full features bundle', '08/04'),
        ('#156', 'feature/all/sprint7-bugfix', 'Fix SOS 400 error + Mapbox blank screen', '08/04'),
        ('#160', 'fix/mapbox-build-errors', 'Fix compile errors, load Mapbox token từ .env', '09/04'),
        ('#161', 'fix/mobile/sprint7-bug-fixes', 'Fix timer sync kẹt 0 + Mapbox missing marker', '09/04'),
        ('#162', 'fix/android-androidx-build', 'Restore gradle.properties với AndroidX enabled', '09/04'),
        ('#165', 'fix/mobile/gps-geofence-sprint7', 'Fix GPS + Geofence detection bugs', '09/04'),
        ('#166', 'fix/mobile/location-history-geofence-events', 'getHistory trả về cả location + geofence events', '09/04'),
        ('#168', 'feature/mobile/location-history-map', 'Thêm map + polyline vào LocationHistoryScreen', '09/04'),
        ('#169', 'fix/mobile/location-history-api-fix', 'Fix LineLayer.lineColorInt và cameraForCoordinateBounds', '10/04'),
        ('#172', 'fix/mobile/geofence-map-bugs', 'Fix biên dịch location_history + race condition', '10/04'),
        ('#173', 'fix/sprint7-geofence-ui-and-history', 'Equal-size dialog buttons + TC-12 live refresh', '10/04'),
        ('#174', 'fix/sprint7-sos-and-countdown', 'SOS confirm dialog, FCM nav, timestamp, marker, history', '10/04'),
        ('#176', 'fix/sprint7-sos-and-countdown', 'Prefix mapbox import, resolve Size conflict', '10/04'),
        ('#179', 'fix/mobile/sprint7-round2-bugs', 'Round 2: geofence dialog wipe + SOS history 404', '10/04'),
        ('#180', 'fix/mobile/global-geofence-sos-events', 'Route geofenceEvent+sosAlert qua SocketService', '11/04'),
        ('#181', 'fix/sprint7-notification-and-sos', 'Fix SOS dialog logic, mic permission, flutter_local_notifications', '11/04'),
        ('#182', 'fix/sprint7-notification-and-sos', 'Enable coreLibraryDesugaring', '11/04'),
        ('#183', 'fix/sos-navigation-and-lints', 'Fix SOS notification navigation fail + dart fixes', '11/04'),
        ('#184', 'fix/sos-navigation-and-lints', 'Fix compiler syntax error in main.dart', '11/04'),
        ('#185', 'fix/sos-fcm-stale-tokens-and-status-badge', 'Auto-clean stale FCM tokens + acknowledge/resolve status', '11/04'),
        ('#186', 'fix/mobile/sos-dialog-double-ui-and-child-guard', 'Prevent SOS dialog trên child app + double UI', '11/04'),
    ]
    for i, row_data in enumerate(pr_mobile):
        add_table_row(tbl4, row_data, bg='FFF2CC' if i % 2 == 0 else None)

    doc.add_paragraph()

    # ── 4. QUÁ TRÌNH KIỂM THỬ ────────────────────────────────────
    add_heading(doc, 'IV. QUÁ TRÌNH KIỂM THỬ VÀ XỬ LÝ LỖI', level=1, color=(31, 73, 125))

    add_para(doc,
        'Sau khi hoàn thành code các tính năng chính vào ngày 08/04 (tất cả được '
        'merge vào develop), nhóm tiến hành test toàn bộ 24 test case theo kế hoạch. '
        'Quá trình test kéo dài từ 09/04 đến 11/04 và phát hiện 6 bug, tất cả đã '
        'được fix trước khi kết thúc sprint.',
        first_line=True)

    add_heading(doc, '4.1. Danh sách Bug phát hiện và xử lý', level=2, size=13)
    tbl5 = doc.add_table(rows=1, cols=6)
    tbl5.style = 'Table Grid'
    make_header_row(tbl5, ['#', 'TC', 'Mô tả lỗi', 'Mức độ', 'Người fix', 'Status'])
    bugs = [
        ('1', 'TC-01', 'Marker vị trí không hiển thị trên bản đồ (thiếu icon asset)', 'High', 'Arix', 'RESOLVED'),
        ('2', 'N/A', 'Đếm ngược thời gian kẹt tại 0 khi phụ huynh cập nhật time limit (delta logic)', 'Critical', 'Arix', 'RESOLVED'),
        ('3', 'TC-06', 'UI Dialog "Lưu vùng an toàn": nút Hủy nhỏ hơn nút Lưu', 'Medium', 'Arix', 'RESOLVED'),
        ('4', 'TC-08', 'UI Dialog "Xóa vùng an toàn": 2 nút không bằng nhau (OverflowBar)', 'Medium', 'Arix', 'RESOLVED'),
        ('5', 'TC-12', 'Màn hình Lịch sử không tự refresh Geofence Event real-time', 'High', 'Arix/Khanh', 'RESOLVED'),
        ('6', 'TC-24', 'POST /api/child/sos ném 500 khi file sai định dạng (Multer error không catch)', 'High', 'Khanh', 'RESOLVED'),
    ]
    severity_colors = {'High': 'FCE4D6', 'Critical': 'FF0000', 'Medium': 'FFEB9C'}
    for bug_data in bugs:
        r = add_table_row(tbl5, bug_data)
        sev = bug_data[3]
        color = severity_colors.get(sev, 'FFFFFF')
        set_cell_bg(r.cells[3], color)
        if sev == 'Critical':
            for run in r.cells[3].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    doc.add_paragraph()

    add_heading(doc, '4.2. Phân tích các lỗi phức tạp', level=2, size=13)
    add_para(doc,
        'Bug số 2 (Countdown timer kẹt tại 0) là lỗi phức tạp nhất sprint: khi phụ '
        'huynh thay đổi time limit qua Socket.IO, delta thời gian bị tính toán sai '
        'do dùng giá trị cũ làm baseline. Fix bằng cách reset timer hoàn toàn thay '
        'vì cộng delta.',
        first_line=True)

    add_para(doc,
        'Bug số 6 (Multer 500 error) xảy ra vì Express không tự bắt lỗi từ Multer '
        'middleware — cần wrap trong try-catch và kiểm tra instanceof MulterError '
        'để trả về HTTP 400 thay vì để lỗi propagate thành 500.',
        first_line=True)

    add_para(doc,
        'Build error Mapbox (prefix import) là vấn đề điển hình khi tích hợp SDK '
        'native vào Flutter — class Size tồn tại ở cả hai namespace. Giải pháp: '
        "dùng `import 'package:mapbox_maps_flutter/mapbox_maps_flutter.dart' as mapbox;` "
        'và gọi đầy đủ mapbox.Size.',
        first_line=True)

    # ── 5. METRICS ────────────────────────────────────────────────
    add_heading(doc, 'V. METRICS SPRINT 7', level=1, color=(31, 73, 125))

    tbl6 = doc.add_table(rows=1, cols=3)
    tbl6.style = 'Table Grid'
    make_header_row(tbl6, ['Chỉ số', 'Giá trị', 'Ghi chú'])
    metrics = [
        ('Tổng commit', '107 commits', 'Bao gồm merge commits'),
        ('Commit thực (non-merge)', '~60 commits', ''),
        ('Pull Requests', '26 PR merged', 'PR #163 đến #188'),
        ('Thời gian sprint', '6 ngày', '07/04 – 12/04/2026'),
        ('Test cases thực hiện', '24 TC', 'Toàn bộ sprint 7 test plan'),
        ('Test cases pass', '24/24 (100%)', 'Không có TC fail còn tồn đọng'),
        ('Bug phát hiện', '6 bugs', 'Từ tester + tự phát hiện'),
        ('Bug resolved', '6/6 (100%)', 'Tất cả đã fix trước khi đóng sprint'),
        ('Models mới (Prisma)', '4 models', 'LocationLog, Geofence, GeofenceEvent, SOSAlert'),
        ('API endpoints mới', '11 endpoints', 'Location (3) + Geofence (5) + SOS (3)'),
        ('Socket.IO events mới', '3 events', 'locationUpdated, geofenceEvent, sosAlert'),
        ('Màn hình mới (Mobile)', '5 màn hình', 'Map, Geofence, LocationHistory, SOS Dialog, SOS History'),
    ]
    for i, m in enumerate(metrics):
        add_table_row(tbl6, m, bg='EBF3FB' if i % 2 == 0 else None)

    doc.add_paragraph()

    # ── 6. ĐÁNH GIÁ & KẾ HOẠCH ──────────────────────────────────
    add_heading(doc, 'VI. ĐÁNH GIÁ VÀ KẾ HOẠCH TIẾP THEO', level=1, color=(31, 73, 125))

    add_heading(doc, '6.1. Điểm mạnh Sprint 7', level=2, size=13)
    add_bullet(doc, 'Hoàn thành 100% kế hoạch trong đúng 1 tuần: tất cả 24 TC pass, không có task backlog.')
    add_bullet(doc, 'Git workflow nghiêm túc: mỗi task = 1 feature branch, mỗi branch = 1 PR, không push trực tiếp vào develop.')
    add_bullet(doc, 'Test coverage tốt: có kế hoạch test chi tiết 24 TC với các edge case (audio upload validation, geofence toggle, SOS không có microphone).')
    add_bullet(doc, 'Code chất lượng: Haversine formula đúng, Socket.IO real-time < 2s latency, Multer error handling an toàn.')

    add_heading(doc, '6.2. Điểm cần cải thiện', level=2, size=13)
    add_bullet(doc, 'Mapbox build issues tốn nhiều thời gian debug (import conflict, blank map, token management). Nên đọc kỹ migration guide trước.')
    add_bullet(doc, 'Số lượng fix PR nhiều (17/26 PR là fix), cho thấy cần test kỹ hơn trước khi merge tính năng.')
    add_bullet(doc, 'Geofence in-memory state cache sẽ mất khi restart server — chấp nhận cho đồ án, cần Redis cho production.')

    add_heading(doc, '6.3. Kế hoạch Sprint 8 (14/04 – 19/04/2026)', level=2, size=13)
    add_para(doc,
        'Sprint 8 tập trung vào Web Filtering, School Mode và Per-app Time Limit:',
        first_line=True)
    add_bullet(doc, 'Backend (Khanh): AppTimeLimit model + API, SchoolSchedule + AllowedSchoolApp models, WebCategory + BlockedCategory models, Custom URL blacklist/whitelist API, School Mode logic.')
    add_bullet(doc, 'Mobile (Arix): Android Native (Kotlin) VpnService cho web filtering, Parent App Per-app Time Limit UI, School Mode UI, Web Filtering UI, Child App áp dụng per-app limit và School Mode.')

    add_heading(doc, '6.4. Tiến độ tổng thể dự án', level=2, size=13)
    add_para(doc,
        'Sau Sprint 7, nhóm đã hoàn thành 7/10 sprint theo kế hoạch. Tất cả tính '
        'năng P0 (phần cốt lõi demo giữa kỳ) và hầu hết P1 (GPS/Geofence/SOS) đã '
        'sẵn sàng. Còn 3 sprint để hoàn thiện Web Filtering, Reports/AI Analysis, '
        'và polish trước khi bảo vệ hội đồng.',
        first_line=True)

    # Progress bar table
    tbl7 = doc.add_table(rows=1, cols=3)
    tbl7.style = 'Table Grid'
    make_header_row(tbl7, ['Sprint', 'Nội dung', 'Trạng thái'])
    sprints = [
        ('Sprint 1', 'Nền tảng & Khởi động', '✅ Hoàn thành'),
        ('Sprint 2', 'Auth & Profile Management', '✅ Hoàn thành'),
        ('Sprint 3', 'Device Management & Socket.IO', '✅ Hoàn thành'),
        ('Sprint 4', 'Time Management & Soft Warning ★', '✅ Hoàn thành'),
        ('Sprint 5', 'Native Android & Lock Screen', '✅ Hoàn thành'),
        ('Sprint 6', 'Demo Giữa Kỳ ★ CHECKPOINT', '✅ Hoàn thành'),
        ('Sprint 7', 'GPS, Geofencing & SOS', '✅ Hoàn thành (tuần này)'),
        ('Sprint 8', 'Web Filter, School Mode, Per-app', '⏳ Sắp bắt đầu'),
        ('Sprint 9', 'Reports, AI & Monitoring', '🔜 Chờ'),
        ('Sprint 10', 'Polish & Bảo vệ hội đồng', '🔜 Chờ'),
    ]
    colors = {
        '✅ Hoàn thành': 'C6EFCE',
        '✅ Hoàn thành (tuần này)': '92D050',
        '⏳ Sắp bắt đầu': 'FFEB9C',
        '🔜 Chờ': 'D9D9D9',
    }
    for sp in sprints:
        r = add_table_row(tbl7, sp)
        set_cell_bg(r.cells[2], colors.get(sp[2], 'FFFFFF'))

    doc.add_paragraph()

    # Ký tên
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('TP. Hồ Chí Minh, ngày 12 tháng 04 năm 2026')
    set_font(run, size=12, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Nhóm 60 — KidFun V3')
    set_font(run, size=12, bold=True)

    # Save
    out_path = '/home/khanh/kidfun-v2/scratch/BaoCao_Sprint7_ChiTiet.docx'
    doc.save(out_path)
    print(f'✅ Saved: {out_path}')
    return out_path


# ═══════════════════════════════════════════════════════════════════
# FILE 2: PHIẾU THEO DÕI TIẾN ĐỘ THỰC HIỆN ĐỒ ÁN CƠ SỞ
# ═══════════════════════════════════════════════════════════════════

def create_progress_form():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HCM (HUTECH)')
    set_font(run, size=12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
    set_font(run, size=12, bold=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PHIẾU THEO DÕI TIẾN ĐỘ THỰC HIỆN ĐỒ ÁN CƠ SỞ')
    set_font(run, size=16, bold=True, color=(31, 73, 125))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('(Báo cáo Tuần 7 — Sprint 7)')
    set_font(run, size=13, bold=False, italic=True)

    doc.add_paragraph()

    # ── THÔNG TIN CHUNG ──────────────────────────────────────────
    add_heading(doc, 'I. THÔNG TIN CHUNG', level=1, size=13, color=(31, 73, 125))

    tbl = doc.add_table(rows=8, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_rows = [
        ('Tên đề tài:', 'KidFun V3 — Ứng dụng kiểm soát thời gian sử dụng thiết bị của trẻ em', 'Chuyên ngành:', 'Công Nghệ Phần Mềm'),
        ('Nhóm:', 'Nhóm 60', 'Số thành viên:', '2'),
        ('Thành viên 1:', 'Khanh — Backend Developer', 'Thành viên 2:', 'Arix — Mobile Developer'),
        ('GVHD:', '', 'Mã lớp:', ''),
        ('Tuần báo cáo:', 'Tuần 7 / 10', 'Sprint:', 'Sprint 7 — GPS, Geofencing & SOS'),
        ('Thời gian:', '07/04/2026 – 12/04/2026', 'Ngày nộp:', '12/04/2026'),
        ('Server:', 'Railway.app (PostgreSQL + Node.js)', 'Repo:', 'GitHub — Khanh-4/kidfun-v2'),
        ('Trạng thái tổng thể:', '✅ Hoàn thành 100% Sprint 7', 'Tiến độ:', '7/10 Sprint hoàn thành'),
    ]

    for i, (k1, v1, k2, v2) in enumerate(info_rows):
        row = tbl.rows[i]
        for j, (txt, is_key) in enumerate([(k1, True), (v1, False), (k2, True), (v2, False)]):
            cell = row.cells[j]
            p2 = cell.paragraphs[0]
            run = p2.add_run(str(txt))
            set_font(run, size=11, bold=is_key)
            if is_key:
                set_cell_bg(cell, 'D9E2F3')

    doc.add_paragraph()

    # ── CÔNG VIỆC ĐÃ THỰC HIỆN ───────────────────────────────────
    add_heading(doc, 'II. CÔNG VIỆC ĐÃ THỰC HIỆN TRONG TUẦN', level=1, size=13, color=(31, 73, 125))

    tbl2 = doc.add_table(rows=1, cols=5)
    tbl2.style = 'Table Grid'
    make_header_row(tbl2, ['STT', 'Thành viên', 'Công việc thực hiện', 'Kết quả', 'Ghi chú'])

    works = [
        ('1', 'Khanh\n(Backend)', 'Thiết kế và migrate 4 Prisma model mới:\nLocationLog, Geofence, GeofenceEvent, SOSAlert', '✅ Hoàn thành', 'Railway PostgreSQL'),
        ('2', 'Khanh', 'Xây dựng Location Tracking API:\n- POST /api/child/location\n- GET location/current\n- GET location/history', '✅ Hoàn thành', 'Socket.IO real-time'),
        ('3', 'Khanh', 'Xây dựng Geofence CRUD API:\n- GET/POST/PUT/DELETE geofences\n- GET geofence events\n- ENTER/EXIT detection (Haversine)', '✅ Hoàn thành', 'In-memory state cache'),
        ('4', 'Khanh', 'Xây dựng SOS Alert API:\n- POST /api/child/sos (multipart + audio)\n- acknowledge / resolve endpoints', '✅ Hoàn thành', 'Multer 5MB limit'),
        ('5', 'Khanh', 'Tích hợp FCM Push Notification:\n- Geofence ENTER/EXIT notification\n- SOS critical priority notification\n- Auto-clean stale tokens', '✅ Hoàn thành', 'Priority MAX cho SOS'),
        ('6', 'Khanh', 'Fix bugs:\n- TC-24: Multer error → 400 thay vì 500\n- Duplicate SOS notification\n- Geofence tên mặc định "Khu vực"', '✅ Hoàn thành', '3 bugs fixed'),
        ('7', 'Arix\n(Mobile)', 'Setup Mapbox SDK:\n- Maven repository config\n- Token từ .env\n- Fix Size import conflict\n- Fix blank map', '✅ Hoàn thành', 'mapbox_maps_flutter'),
        ('8', 'Arix', 'GPS Tracking Service:\n- Foreground: 30s interval\n- Background: 5min interval\n- WidgetsBindingObserver lifecycle', '✅ Hoàn thành', 'geolocator package'),
        ('9', 'Arix', 'Parent App Map Screen:\n- Real-time marker (Socket.IO)\n- Fetch current location\n- FAB refresh', '✅ Hoàn thành', 'PointAnnotationManager'),
        ('10', 'Arix', 'Parent App Geofence UI:\n- Tap to create (marker + circle)\n- Radius slider (50–5000m)\n- Delete by tap polygon\n- Real-time ENTER/EXIT dialog', '✅ Hoàn thành', 'CircleAnnotationManager'),
        ('11', 'Arix', 'Parent App Location History:\n- DatePicker + API\n- Polyline vẽ đường đi\n- Kết hợp geofence events', '✅ Hoàn thành', 'LineLayer + GeoJsonSource'),
        ('12', 'Arix', 'Child App SOS Button:\n- Confirm dialog\n- Fast alert < 2s\n- Ghi âm 15s + upload', '✅ Hoàn thành', 'record package'),
        ('13', 'Arix', 'Parent App SOS Alert:\n- Dialog barrierDismissible=false\n- Nghe audio / Xem vị trí / Gọi con\n- SOS History + status badge', '✅ Hoàn thành', 'flutter_local_notifications'),
        ('14', 'Arix', 'Fix bugs Mobile:\n- Timer countdown kẹt 0 (critical)\n- Mapbox blank screen + marker\n- SOS double dialog\n- FCM navigation\n- Line color API', '✅ Hoàn thành', '~10 bugs fixed'),
        ('15', 'Cả nhóm', 'Kiểm thử Sprint 7:\n- Thực hiện 24 test case\n- Ghi nhận 6 bugs\n- Fix tất cả bugs', '✅ 24/24 Pass', '100% pass rate'),
    ]

    for i, w in enumerate(works):
        r = add_table_row(tbl2, w, bg='EBF3FB' if i % 2 == 0 else None)
        # Justify long text cells
        for cell in r.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # ── KẾT QUẢ KIỂM THỬ ────────────────────────────────────────
    add_heading(doc, 'III. KẾT QUẢ KIỂM THỬ', level=1, size=13, color=(31, 73, 125))

    tbl3 = doc.add_table(rows=1, cols=5)
    tbl3.style = 'Table Grid'
    make_header_row(tbl3, ['Module', 'Tổng TC', 'Pass', 'Fail', 'Kết quả'])
    test_results = [
        ('GPS Tracking (TC-01~05)', '5', '5', '0', '✅ PASS'),
        ('Geofence (TC-06~12)', '7', '7', '0', '✅ PASS'),
        ('SOS (TC-13~21)', '9', '9', '0', '✅ PASS'),
        ('API Validation (TC-22~24)', '3', '3', '0', '✅ PASS'),
        ('TỔNG', '24', '24', '0', '✅ 100%'),
    ]
    for i, tr in enumerate(test_results):
        r = add_table_row(tbl3, tr, bg='C6EFCE' if i == 4 else ('EBF3FB' if i % 2 == 0 else None))
        if i == 4:
            for cell in r.cells:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    doc.add_paragraph()

    # ── TIẾN ĐỘ SO VỚI KẾ HOẠCH ─────────────────────────────────
    add_heading(doc, 'IV. TIẾN ĐỘ SO VỚI KẾ HOẠCH', level=1, size=13, color=(31, 73, 125))

    tbl4 = doc.add_table(rows=1, cols=4)
    tbl4.style = 'Table Grid'
    make_header_row(tbl4, ['Hạng mục', 'Kế hoạch', 'Thực hiện', 'So sánh'])
    progress = [
        ('GPS Tracking', 'Hoàn thành Sprint 7', 'Hoàn thành 08/04', '✅ Đúng kế hoạch'),
        ('Geofencing CRUD', 'Hoàn thành Sprint 7', 'Hoàn thành 08/04', '✅ Đúng kế hoạch'),
        ('Geofence Detection', 'Hoàn thành Sprint 7', 'Hoàn thành 08/04', '✅ Đúng kế hoạch'),
        ('SOS Button + Alert', 'Hoàn thành Sprint 7', 'Hoàn thành 08/04', '✅ Đúng kế hoạch'),
        ('Push Notification (Geofence+SOS)', 'Hoàn thành Sprint 7', 'Hoàn thành 08/04', '✅ Đúng kế hoạch'),
        ('Mapbox SDK Integration', 'Hoàn thành Sprint 7', 'Hoàn thành (nhiều vòng fix)', '⚠️ Tốn nhiều thời gian'),
        ('Test + Bug Fix', 'Kết thúc 12/04', 'Kết thúc 11/04', '✅ Sớm hơn 1 ngày'),
        ('Sprint 7 tổng thể', '100% hoàn thành', '100% hoàn thành', '✅ Đúng kế hoạch'),
    ]
    for i, p_data in enumerate(progress):
        r = add_table_row(tbl4, p_data, bg='EBF3FB' if i % 2 == 0 else None)
        status = p_data[3]
        if '✅' in status:
            set_cell_bg(r.cells[3], 'C6EFCE')
        elif '⚠️' in status:
            set_cell_bg(r.cells[3], 'FFEB9C')

    doc.add_paragraph()

    # ── ĐÁNH GIÁ ─────────────────────────────────────────────────
    add_heading(doc, 'V. ĐÁNH GIÁ CỦA NHÓM', level=1, size=13, color=(31, 73, 125))

    tbl5 = doc.add_table(rows=4, cols=2)
    tbl5.style = 'Table Grid'
    eval_data = [
        ('Mức độ hoàn thành công việc:', '100% — Tất cả task trong Sprint 7 đã hoàn thành, 24/24 TC pass.'),
        ('Khó khăn gặp phải:', 'Tích hợp Mapbox mất nhiều thời gian (namespace conflict, blank map, token management). Nhiều vòng fix build error liên quan đến Gradle và Android.'),
        ('Hướng giải quyết:', 'Đọc kỹ documentation, tách biệt token khỏi VCS bằng .env, prefix import để giải quyết namespace conflict.'),
        ('Kế hoạch tuần tới (Sprint 8):', 'Bắt đầu 14/04. Backend: Per-app Time Limit API, School Mode, Web Filtering models. Mobile: VpnService (web filter), Per-app UI, School Mode UI.'),
    ]
    for i, (k, v) in enumerate(eval_data):
        row = tbl5.rows[i]
        p2 = row.cells[0].paragraphs[0]
        run = p2.add_run(k)
        set_font(run, size=11, bold=True)
        set_cell_bg(row.cells[0], 'D9E2F3')
        row.cells[0].width = Cm(5)

        p3 = row.cells[1].paragraphs[0]
        run = p3.add_run(v)
        set_font(run, size=11)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()

    # ── METRICS ──────────────────────────────────────────────────
    add_heading(doc, 'VI. CHỈ SỐ KỸ THUẬT', level=1, size=13, color=(31, 73, 125))

    tbl6 = doc.add_table(rows=1, cols=4)
    tbl6.style = 'Table Grid'
    make_header_row(tbl6, ['Chỉ số', 'Khanh (Backend)', 'Arix (Mobile)', 'Tổng'])
    metric_rows = [
        ('Commits', '~25', '~35', '107 (kể cả merge)'),
        ('Pull Requests', '8 PR', '17 PR', '26 PR'),
        ('Files thay đổi', 'Backend + Prisma', 'Mobile + Android', 'Toàn bộ dự án'),
        ('API endpoints mới', '11 endpoints', 'N/A', '11 endpoints'),
        ('Màn hình mới', 'N/A', '5 màn hình', '5 màn hình'),
        ('Bug đã fix', '3 bugs', '~10 bugs', '6 bugs chính thức'),
    ]
    for i, mr in enumerate(metric_rows):
        add_table_row(tbl6, mr, bg='EBF3FB' if i % 2 == 0 else None)

    doc.add_paragraph()

    # ── XÁC NHẬN ─────────────────────────────────────────────────
    add_heading(doc, 'VII. XÁC NHẬN', level=1, size=13, color=(31, 73, 125))

    tbl7 = doc.add_table(rows=4, cols=3)
    tbl7.style = 'Table Grid'
    sign_headers = ['', 'Nhóm trưởng', 'GVHD']
    r0 = tbl7.rows[0]
    for i, h in enumerate(sign_headers):
        p2 = r0.cells[i].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p2.add_run(h)
        set_font(run, size=11, bold=True)
        set_cell_bg(r0.cells[i], '4472C4')
        run.font.color.rgb = RGBColor(255, 255, 255)

    labels = ['Họ tên:', 'Chữ ký:', 'Ngày xác nhận:']
    values = [
        ('Khanh', ''),
        ('', ''),
        ('12/04/2026', ''),
    ]
    for i, (label, (v1, v2)) in enumerate(zip(labels, values)):
        row = tbl7.rows[i + 1]
        p2 = row.cells[0].paragraphs[0]
        run = p2.add_run(label)
        set_font(run, size=11, bold=True)
        set_cell_bg(row.cells[0], 'D9E2F3')

        p3 = row.cells[1].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p3.add_run(v1)
        set_font(run, size=11)

        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TP. Hồ Chí Minh, ngày 12 tháng 04 năm 2026')
    set_font(run, size=11, italic=True)

    out_path = '/home/khanh/kidfun-v2/scratch/PhieuTheoDoi_TienDo_Tuan7.docx'
    doc.save(out_path)
    print(f'✅ Saved: {out_path}')
    return out_path


if __name__ == '__main__':
    p1 = create_detailed_report()
    p2 = create_progress_form()
    print(f'\nDone! Files saved:\n  1. {p1}\n  2. {p2}')
